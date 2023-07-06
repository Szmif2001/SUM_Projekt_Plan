from docx import Document
import aspose.words as aw
import csv
import mechanize
from time import sleep
import requests
import os
import shutil

def download_link(l):
    filename = ''
    url = l.url
    if url.find('/'):
        filename = url.rsplit('/', 1)[1]

    try:
        r = requests.get(url, allow_redirects = True)
        open(filename, mode = 'wb').write(r.content)
    except:
        pass #BUG naprawić 

def move_file(filename):
    destination = 'Dump'
    if filename.endswith('.docx') or filename.endswith('.doc'):
        if filename.startswith('CMUS'):
            destination = 'Wydzialy\\WNOZK\\Coaching_medyczny'
        if filename.startswith('ER'):
            destination = 'Wydzialy\\WNOZK\\Elektroradiologia'
        if filename.startswith('Fizjoterapia'):
            destination = 'Wydzialy\\WNOZK\\Fizjoterapia'
        if filename.startswith('PLL') or filename.startswith('PLUS') or filename.startswith('PLUN'):
            destination = 'Wydzialy\\WNOZK\\Pielegniarstwo'
        if filename.startswith('POL') or filename.startswith('POUS') or filename.startswith('POUN'):
            destination = 'Wydzialy\\WNOZK\\Poloznictwo'
        shutil.move(filename, destination)

def website(url):
    browser = mechanize.Browser()
    browser.open(url)

    filetypes = ['.doc', '.docx']
    files = []

    for l in browser.links():
        for t in filetypes:
            if t in str(l):
                files.append(l)
    
    browser.close()

    for l in files:
        download_link(l)
        sleep(1)

    files_in_dir = os.listdir()
    for i in files_in_dir:
        move_file(i)

def doc_del(file):
        if file.endswith('.doc'):
            os.remove(file)

def doc_docx(file):
    if file.endswith('.doc'):
        Doc = aw.Document(file)
        Doc.save(file.rsplit('.', 1)[0] + '.docx')
        file = file.rsplit('.', 1)[0] + '.docx'

    Docx = Document(file)
    out_temp = []
    out = []

    for table in Docx.tables:
        for row in table.rows:
            temp = []
            for cell in row.cells:
                temp.append(cell.text.encode('utf8'))
            out_temp.append(temp)

    dni = out_temp[0]
    zaj = out_temp[1]

    for i in range(0, min(len(dni), len(zaj)), 1):
        strings = zaj[i].decode('utf8').split('\n')
        temp = [dni[i].decode('utf8'), strings]
        out.append(temp)

    with open(file.rsplit('.', 1)[0] + '.csv', mode = 'w', encoding = 'cp1250') as csv_file:
        writer = csv.writer(csv_file)
        writer.writerow(r for r in out)

# BUG: Poszukać innego sposobu encodowania lub ręcznie tłumaczyć dziwne znaczki na polskie litery


def xls_xlsx(file):
    print('to be implemented')



url = ['https://student.sum.edu.pl/dziekanat-wydzialu-nauk-o-zdrowiu-w-katowicach/#harmonogramy', 'https://student.sum.edu.pl/dziekanat-wydzialu-lekarskiego-w-katowicach/#efekty-ksztacenia-i-harmonogramy', 'https://student.sum.edu.pl/dziekanat-wydzialu-lekarskiego-z-oddzialem-lekarsko-dentystycznym-w-zabrzu/#harmonogramy', 'https://student.sum.edu.pl/dziekanat-wydzialu-farmaceutycznego-z-oddzialem-medycyny-laboratoryjnej-w-sosnowcu/#harmonogramy', 'https://student.sum.edu.pl/dziekanat-wydzialu-zdrowia-publicznego-w-bytomiu/#harmonogramy']
#for i in url:
#    website(i)

wydzialy = os.listdir('Wydzialy')
wydzialy.pop(0)

#doc_docx('Wydzialy\\WNOZK\\Fizjoterapia\\Fizjoterapia-jmgr-1r-l.doc')

for w in wydzialy:
    kierunki = os.listdir(f'Wydzialy\\{w}')
    for k in kierunki:
        files = os.listdir(f'Wydzialy\\{w}\\{k}')
        for f in files:
            if f.endswith('.doc') or f.endswith('.docx'):
                doc_docx(f'Wydzialy\\{w}\\{k}\\{f}')
            if f.endswith('.xls') or f.endswith('.xlsx'):
                xls_xlsx(f'Wydzialy\\{w}\\{k}\\{f}')

#TODO: osobny plik.py dla kazdego wydzialu
